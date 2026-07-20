# -*- coding: utf-8 -*-
"""Traceability audit: every node / edge / KPI of the data-flow figure vs the manuscript."""
import docx, re, sys

sys.stdout.reconfigure(encoding="utf-8", errors="replace")
DOC = r"C:/Users/Administrador/Downloads/PUSMA_tool_wear_prognosis_MDPI.docx"

d = docx.Document(DOC)
paras = [p.text for p in d.paragraphs]
tbl = " ".join(c.text for t in d.tables for r in t.rows for c in r.cells)
FULL = " ".join(paras) + " " + tbl
FLAT = re.sub(r"\s+", " ", FULL)

# section index: map a char position -> nearest preceding heading
heads = []
pos = 0
for p in paras:
    t = p.text if hasattr(p, "text") else p
    if re.match(r"^(\d+(\.\d+)?\.\s|Appendix|Abstract:|Keywords:)", t.strip()) and len(t) < 90:
        heads.append((pos, t.strip()[:44]))
    pos += len(t) + 1

def where(pat):
    m = re.search(pat, FLAT, re.I)
    if not m:
        return None
    # locate in the paragraph stream for a section estimate
    p = 0
    sect = "—"
    for txt in paras:
        if re.search(pat, txt, re.I):
            for hp, hn in heads:
                if hp <= p:
                    sect = hn
            return sect
        p += len(txt) + 1
    return "(table/caption)"

# (group, node label in figure, regex evidence that it is described in the paper)
CHECKS = [
 # ---- acquisition ----
 ("NODE", "a1 · Hanwha XD26II-V Swiss-type centre", r"Hanwha XD26II-V"),
 ("NODE", "a1 · 18 tools = full 3x3x2 DOE", r"18 tools\s*=\s*full DOE|eighteen-tool|18 distinct cutting conditions"),
 ("NODE", "a1 · DOE LEVELS vc 55/70/80", r"55/70/80|55,\s*70,\s*80|v_?c\s*∈"),
 ("NODE", "a1 · DOE LEVELS f 0.08/0.20/0.30", r"0\.08/0\.20/0\.30|0\.08,\s*0\.20,\s*0\.30"),
 ("NODE", "a1 · ONE tool per recipe", r"one tool per condition|single tool|no replicate tools"),
 ("NODE", "a2 · axial + radial accelerometers", r"axial \(A\) and radial \(R\)"),
 ("NODE", "a2 · one burst per contact, 6 per cycle", r"six consecutive contacts|one energy burst per"),
 ("NODE", "a2 · AE recorded, set aside", r"acoustic emission, though recorded, is not pursued"),
 ("NODE", "a3 · Keyence VHX-7000", r"Keyence VHX-7000"),
 ("NODE", "a3 · VB ex situ at cycle stops", r"ex situ|measured .{0,30}only at the end of each cycle"),
 ("NODE", "a3 · ~5 inspections per tool", r"five inspection points|≈5 VB measurements|roughly five labels"),
 # ---- raw stores / preparation ----
 ("NODE", "b1 · raw signal store (per-contact bursts)", r"segmented per contact"),
 ("NODE", "b1 · TXT file format", r"\bTXT\b"),
 ("NODE", "b2 · VB target table / 172 records", r"\b172\b"),
 ("NODE", "b2 · 18 chipping values 127-291 um", r"127 to 291|127–291"),
 ("NODE", "c1 · signal QA + segmentation", r"segmented per contact|signal quality"),
 ("NODE", "c2 · 294 features per cut", r"294 (?:generic )?(?:vibration )?(?:features|descriptors)"),
 ("NODE", "c2 · RMS / energies / kurtosis / crest / freq", r"root-mean-square level, signal and band energies"),
 ("NODE", "c2 · wavelet-band energies", r"wavelet-band energies"),
 ("NODE", "c3 · features_experiment.csv (filename)", r"features_experiment\.csv"),
 ("NODE", "c3 · 296 numeric columns", r"296-descriptor|296 "),
 # ---- the wall ----
 ("NODE", "WALL · LOOCV, 18 folds, one tool held out", r"one tool is held out entirely"),
 ("NODE", "WALL · only first m readings visible", r"only its first m measurements"),
 ("NODE", "WALL · future sealed for scoring", r"sealed and used exclusively for scoring"),
 ("NODE", "WALL · scaling/selection/tuning/calib in-fold", r"all scaling is fitted on training|nested inside the training folds"),
 # ---- sensor lane ----
 ("NODE", "cur · curation 294 -> 181", r"curated to 181"),
 ("NODE", "s1 · consensus feature selection", r"consensus feature selection"),
 ("NODE", "s2 · fold-safe augmentation", r"augmentation"),
 ("NODE", "s3 · nested hyper-parameter tuning", r"nested hyper-parameter tuning"),
 ("NODE", "s4 · A/R fusion, 6 strategies", r"six standard strategies|six fusion"),
 ("NODE", "s5 · ridge / LASSO / elastic net", r"ridge regression \[\d+\], the LASSO"),
 ("NODE", "s5 · random forest", r"random forest"),
 ("NODE", "s5 · XGBoost", r"XGBoost"),
 ("NODE", "s5 · MLP", r"multilayer perceptron"),
 ("NODE", "s5 · PLS", r"partial least squares"),
 ("NODE", "s5 · physics-equation hybrids (eq. 11)", r"residual learning|physics-regressor variant"),
 ("NODE", "sout · cross-tool reading fails, R2<0", r"R² < 0|stays on the wrong side of zero"),
 ("NODE", "sout · prognosability 0.04", r"prognosability across tools, however, is 0\.04|prognosability .{0,20}0\.04"),
 # ---- physics lane ----
 ("NODE", "p1 · fleet exponent, pooled SSE, 17 tools", r"trained on the other 17 tools|pooled residual over all training"),
 ("NODE", "p1 · exponent grid p in [0.20, 1.00]", r"p ∈ \[0\.20, 1\.00\]"),
 ("NODE", "p1 · p* ~ 0.20", r"p ≈ 0\.20|stable at p ≈ 0\.20"),
 ("NODE", "p2 · tau = t^p linearising coordinate", r"linearising coordinate"),
 ("NODE", "p2 · Theil-Sen / Siegel", r"Theil–Sen.{0,40}Siegel|repeated-median"),
 ("NODE", "p2 · extrapolation weights w ~ tau^gamma, gamma=3", r"γ = 3"),
 ("NODE", "p2 · gross-error guard at 5 sigma_meas", r"5·σ_meas|5\s*·\s*σ"),
 ("NODE", "p3 · VB(t) = b + a t^p", r"VB\(t\) = b|monotone power law"),
 ("NODE", "p3 · monotone + decelerating (eq. 3)", r"monotone growth .{0,30}deceleration|non-decreasing and concave"),
 ("NODE", "p3 · uses VB readings ONLY (no sensors)", r"only signal that demonstrably generalises|from VB readings"),
 # ---- calibration / decision ----
 ("NODE", "k1 · conformal band, out-of-fold residuals", r"LOOCV residuals as a calibration set"),
 ("NODE", "k1 · ~118 pooled residuals", r"118 pooled residuals"),
 ("NODE", "k1 · Mondrian horizon bins <=1 / 2-3 / >=4", r"near \(≤ 1 step ahead\), mid \(2–3\) and far \(≥ 4\)"),
 ("NODE", "k1 · +-19 um near / 52.5 um mean / 90.1%", r"±19 µm"),
 ("NODE", "k2 · logit hazard on 172 cycles + 18 events", r"172 cycle records and their 18 terminal events"),
 ("NODE", "k2 · h<=0.10 -> VB_safe ~ 167 um", r"VB_safe ≈ 167"),
 ("NODE", "k2 · P10 cross-check 135 um", r"135 µm"),
 ("NODE", "doe · unreplicated DOE inference (Lenth/ANOVA)", r"Lenth"),
 ("NODE", "doe · cooling moves wear LEVELS", r"cooling: it is the largest effect|do respond to cooling"),
 ("NODE", "d3 · online Kalman, tau-clock, event-triggered", r"event-triggered"),
 ("NODE", "d3 · next cut 3.7-4.0 um", r"3\.7 µm"),
 ("NODE", "d1 · RUL ladder {120,150,175,200}", r"\{120, 150, 175, 200\}"),
 ("NODE", "d1 · 16 events, 94% covered", r"94%"),
 ("NODE", "d2 · stop when UPPER band edge crosses VB_safe", r"UPPER edge of the calibrated band"),
 # ---- edge labels ----
 ("EDGE", "no transferable signal", r"does not generalise|no transferable"),
 ("EDGE", "18 chipping events -> hazard", r"eighteen genuine end-of-life events|18 terminal events"),
 ("EDGE", "radial spectral kurtosis, p = 0.010", r"spectral kurtosis is nominally significant"),
 ("EDGE", "alarm (anomaly check)", r"anomaly (?:detection|check)"),
 # ---- payoff KPIs ----
 ("KPI", "5.6 um = 2.8% of the 200 um budget", r"2\.8% of the 200 µm"),
 ("KPI", "90% band, +-19 um next cycle", r"±19 µm"),
 ("KPI", "167 um safe stop, one-cycle risk <= 10%", r"h_max = 0\.10|risk budget"),
 ("KPI", "94% of 16 real RUL events", r"94% of them|16 .{0,40}events"),
 ("KPI", "0 further readings after the 4th", r"runs measurement-free|no further measurement at all"),
]

print("=" * 96)
print("TRAZABILIDAD: cada elemento del diagrama de flujo -> manuscrito")
print("=" * 96)
missing = []
for kind, label, pat in CHECKS:
    sect = where(pat)
    ok = sect is not None
    if not ok:
        missing.append((kind, label, pat))
    mark = "OK  " if ok else "FALTA"
    print(f"[{mark}] {kind:4s} | {label:52s} | {sect or '--- NO ENCONTRADO ---'}")
print("-" * 96)
print(f"TOTAL: {len(CHECKS)} elementos | presentes: {len(CHECKS)-len(missing)} | FALTANTES: {len(missing)}")
if missing:
    print("\nELEMENTOS DEL FLUJO QUE NO APARECEN EN EL PAPER:")
    for kind, label, pat in missing:
        print(f"   - [{kind}] {label}     (patrón: {pat})")
